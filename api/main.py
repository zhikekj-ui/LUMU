import time
import uuid
import os
import re
import io
import asyncio
from slowapi.errors import RateLimitExceeded
from prometheus_client import Counter, Histogram, generate_latest, CONTENT_TYPE_LATEST
from middleware.rate_limit import limiter, rate_limit_exceeded_handler
from core.logging_config import configure_logging, get_logger

"""FastAPI application — API routes + static file serving."""
import json
from contextlib import asynccontextmanager
from pathlib import Path

from fastapi import FastAPI, HTTPException, Header, Depends, Request, UploadFile, File, Form
from fastapi.responses import HTMLResponse, JSONResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel

from agent.core import Agent
from agent.tracing import get_tracer
from tools.registry import ToolRegistry
from providers.registry import discover_providers
from plugins.loader import PluginLoader
from mcp.bridge import MCPBridge
from scheduler.scheduler import scheduler
import config
from core.user_config import (
    load_config, save_config, get_provider_key, set_provider_key,
    get_tts_config, set_tts_config, get_stt_config,
    get_model_preference, set_model_preference,
    get_system_prompt, set_system_prompt,
    get_embedding_config, set_embedding_config,
    get_enabled_models, set_enabled_models,
)

# --- Bootstrap ---
discover_providers()

tool_registry = ToolRegistry()
tool_registry.discover()

# Load plugins
plugin_loader = PluginLoader()
plugin_loader.discover()
plugin_loader.register_tools(tool_registry)

# MCP bridge
mcp_bridge = MCPBridge()

# Load saved model preference (overrides .env defaults)
_saved_pref = get_model_preference()
_init_provider = _saved_pref.get("provider", config.DEFAULT_PROVIDER)
_init_model = _saved_pref.get("model", config.DEFAULT_MODEL)

agent = Agent(
    provider_name=_init_provider,
    model=_init_model,
    tool_registry=tool_registry,
)

# Wire scheduler to agent
scheduler.set_agent(agent)


# --- Lifespan ---
@asynccontextmanager
async def lifespan(app: FastAPI):
    # Startup: initialize logging
    configure_logging(log_level="INFO", json_format=False)
    logger = get_logger("lifespan")
    logger.info("Agent Framework starting up")
    # Startup: connect MCP servers
    await mcp_bridge.connect_all(tool_registry)
    # Start scheduler
    await scheduler.start()
    # Startup: pre-initialize session/task singletons off the event loop.
    # SessionManager.__init__ does synchronous SQLite schema work + legacy JSON
    # import; doing it here (in a worker thread) avoids constructing it lazily
    # inside the first chat request's async event loop (which would briefly freeze
    # request handling). RLock in session_manager makes _get_tracker/_get_manager reentrant.
    try:
        from agent.session_manager import _get_tracker
        await asyncio.to_thread(_get_tracker)
        logger.info("Session/task manager pre-initialized at startup")
    except Exception as e:
        logger.warning(f"Session/task pre-init failed (non-fatal): {e}")
    # Start channels
    try:
        from channels.router import channel_router
        await channel_router.start_all(agent)
    except Exception as e:
        print(f"[channels] Startup error: {e}")
    yield
    # Shutdown: stop scheduler + disconnect MCP + stop channels
    await scheduler.stop()
    await mcp_bridge.disconnect_all()
    try:
        from channels.router import channel_router
        await channel_router.stop_all()
    except Exception:
        pass


app = FastAPI(title="Agent Framework", version="0.5.0", lifespan=lifespan)
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, rate_limit_exceeded_handler)

# Let plugins register API routes
plugin_loader.register_routes(app)

STATIC_DIR = Path(__file__).parent / "static"


@app.get("/health")
async def health():
    return {
        "status": "ok",
        "provider": getattr(agent, "provider_name", config.DEFAULT_PROVIDER),
        "model": getattr(agent, "model", config.DEFAULT_MODEL),
        "tools_loaded": len(getattr(tool_registry, "_tools", {})),
    }


@app.get("/api/onboarding")
async def onboarding():
    """结构化「出场配置」：供前端首屏欢迎区直接消费（名字/标语/能力分组/示例/使用规矩）。"""
    return {
        "agent": "LUMU",
        "tagline": "常驻你私人服务器的 AI 助理 · 记忆生命体",
        "greeting": "你好，我是 LUMU —— 住在你服务器上的 AI 助理。我会记住你的偏好、越用越懂你。",
        "capability_groups": [
            {"icon": "🧠", "title": "记忆与知识", "desc": "记住偏好/对话/经验，随时召回；从知识库检索资料",
             "examples": ["记住我喜欢的报告风格", "从知识库找一下上次的方案"]},
            {"icon": "🌐", "title": "浏览与检索", "desc": "打开网页、抓取正文、联网查最新信息",
             "examples": ["看看今天 AI 领域有什么热点", "把这篇网页要点整理成笔记"]},
            {"icon": "💻", "title": "执行与文件", "desc": "读写文件、跑代码、管理系统与进程",
             "examples": ["把 data 下的 csv 汇总成图表", "跑一下这个 Python 脚本"]},
            {"icon": "⏰", "title": "定时与自动化", "desc": "设置定时任务、心跳提醒",
             "examples": ["每天早 8 点给我发今日早报", "每周一提醒我写周报"]},
            {"icon": "🤖", "title": "子代理与协作", "desc": "复杂任务拆给多个子代理并行处理",
             "examples": ["调研三个竞品并做对比"]},
            {"icon": "🛡️", "title": "安全护栏", "desc": "危险命令先请求你确认，不擅自执行破坏性操作",
             "examples": []},
            {"icon": "🖼️", "title": "多模态", "desc": "看懂图片/截图，生成图表与可视化",
             "examples": ["这张图讲了什么", "生成一张销售趋势图"]},
            {"icon": "🔍", "title": "深度推理", "desc": "链式/树状思考，解决复杂问题",
             "examples": ["帮我规划一次系统迁移方案"]},
        ],
        "usage_rules": [
            "危险操作会先征求你确认，不会擅自执行",
            "长任务会先给计划，再分步执行",
            "用中文交流就用中文回复",
        ],
        "suggested_prompts": [
            "帮我看看今天 AI 领域有什么热点",
            "把这篇网页的要点整理成笔记",
            "每天早上 8 点给我发一份今日早报",
        ],
    }


@app.get("/metrics")
async def metrics():
    from fastapi import Response
    return Response(content=generate_latest(), media_type=CONTENT_TYPE_LATEST)


# --- Auth ---
logger_auth = get_logger("auth")

async def verify_api_key(request: Request, authorization: str = Header(default="")):
    """访问守卫：本机直连零鉴权；对外暴露时需要一次性访问口令。

    策略实现见 core/access_guard —— 安全跟着「暴露面」走，不跟着「用户身份」走。
    个人智能体不需要账号体系：本机使用完全无感，只有当实例可被外部访问时
    （绑定非环回地址 / 经反向代理 / LUMU_PUBLIC=1）才要求口令。
    """
    from core.access_guard import check as _guard_check, request_is_exposed
    try:
        _guard_check(request, authorization)
    except HTTPException:
        try:
            if request_is_exposed(request):
                logger_auth.warning("unauthorized_request", path=str(request.url.path))
        except Exception:
            pass
        raise


# --- Request models ---
class ChatRequest(BaseModel):
    message: str
    session_id: str | None = None
    images: list[str] | None = None  # v8: base64 or URL images
    files: list[dict] | None = None  # 通用文件附件：[{name, mime, data(base64)}]
    space: str = "work"  # 空间隔离：work / personal
    # —— 对话级模型/参数覆盖（可选，仅对本轮对话临时生效，不持久化、不影响全局配置）——
    provider: str | None = None  # 覆盖供应商（如 anthropic / openai）
    model: str | None = None     # 覆盖具体模型（如 claude-3-5-sonnet-20241022）
    temperature: float | None = None  # 覆盖采样温度


class MemoryRequest(BaseModel):
    key: str
    content: str
    category: str = "general"
    confirmed: bool = False
    importance: float = None
    metadata: dict = None
    store: str = "primary"  # primary=MemoryManager / semantic=SemanticMemory


# --- Routes ---
# ---------------------------------------------------------------- 访问模式（小白开关：本机 / 对外分享）
def _public_base(request: Request) -> str:
    """根据反向代理头拼出当前可访问的基址（小白无需懂这个）。"""
    proto = request.headers.get("x-forwarded-proto") or request.url.scheme or "http"
    host = request.headers.get("x-forwarded-host") or request.headers.get("host") or ""
    if host:
        return f"{proto}://{host}".rstrip("/")
    return str(request.base_url).rstrip("/")


@app.get("/api/access")
async def get_access(request: Request, _=Depends(verify_api_key)):
    from core.access_guard import auth_disabled, request_is_exposed, get_token
    exposed = request_is_exposed(request)
    disabled = auth_disabled()
    token = get_token(create=False)
    base = _public_base(request)
    link = f"{base}/?token={token}" if token else None
    mode = "local" if (disabled or not exposed) else "share"
    return JSONResponse({
        "mode": mode, "exposed": exposed, "auth_disabled": disabled,
        "token_present": bool(token), "share_link": link,
    })


@app.post("/api/access")
async def post_access(request: Request, _=Depends(verify_api_key)):
    from core.access_guard import rotate_token, get_token, request_is_exposed
    try:
        body = await request.json()
    except Exception:
        body = {}
    action = (body.get("action") or "").strip()
    base = _public_base(request)
    if action == "enable":
        token = get_token(create=True)
        return JSONResponse({"ok": True, "mode": "share", "share_link": f"{base}/?token={token}"})
    if action == "rotate":
        token = rotate_token()
        return JSONResponse({"ok": True, "mode": "share", "share_link": f"{base}/?token={token}"})
    if action == "disable":
        if request_is_exposed(request):
            raise HTTPException(status_code=400,
                detail="实例已对外暴露，出于安全不能关闭口令（否则任何人都能直接调用后台）。请保持对外分享模式。")
        return JSONResponse({"ok": True, "mode": "local", "share_link": None})
    raise HTTPException(status_code=400, detail="未知 action：" + action)


@app.post("/api/access/activate")
async def activate_access(request: Request):
    """首次打开的机器：点一下"确认进入"即自动激活本机口令（种 cookie），无需任何链接或终端。"""
    from core.access_guard import get_token
    token = get_token(create=True)
    resp = JSONResponse({"ok": True})
    resp.set_cookie(
        "lumu_token", token,
        max_age=31536000, httponly=True, samesite="lax", path="/",
    )
    return resp


@app.get("/", response_class=HTMLResponse)
async def index():
    html = (STATIC_DIR / "index.html").read_text()
    return HTMLResponse(
        content=html,
        headers={
            "Cache-Control": "no-cache, no-store, must-revalidate",
            "Pragma": "no-cache",
            "Expires": "0",
        },
    )


@limiter.limit("10/minute")
@app.post("/api/chat")
async def chat(req: ChatRequest, request: Request, _=Depends(verify_api_key)):
    try:
        result = await agent.chat(req.message, req.session_id, images=req.images, files=req.files, space=req.space)
        return JSONResponse(result)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@limiter.limit("20/minute")
@app.post("/api/chat/stream")
async def chat_stream(req: ChatRequest, request: Request, _=Depends(verify_api_key)):
    """SSE streaming endpoint — yields tokens as they arrive.
    支持对话级模型/参数覆盖：req.provider / req.model / req.temperature 仅对本轮临时生效，不影响全局 model_preference。
    """

    async def event_generator():
        from tools.file_hub import flush_session_files, _CUR_SESSION
        from providers.registry import get as _get_provider
        from agent.context import ContextEngine
        sid = req.session_id or "__default__"
        _CUR_SESSION.set(sid)
        # —— 对话级模型/参数覆盖（临时切换，整轮流结束后还原）——
        _restore = None
        async with agent._chat_lock:
            if req.provider or req.model or req.temperature is not None:
                try:
                    _orig = (
                        agent.provider_name,
                        agent.provider,
                        agent.model,
                        getattr(agent, "_override_temperature", None),
                        agent.context,
                    )
                    if req.provider and req.provider != agent.provider_name:
                        _p = _get_provider(req.provider)
                        if not _p:
                            raise ValueError(f"供应商 '{req.provider}' 不存在")
                        agent.provider_name = req.provider
                        agent.provider = _p
                        agent.context = ContextEngine(context_window=_p.context_window)
                    if req.model:
                        agent.model = req.model
                    if req.temperature is not None:
                        agent._override_temperature = req.temperature
                    _restore = _orig
                except Exception as _e:
                    # 覆盖失败则回退到全局模型，不阻塞对话
                    yield f"data: {json.dumps({'type': 'error', 'content': '模型切换失败：' + str(_e)}, ensure_ascii=False)}\n\n"
            try:
                async for event in agent.stream_chat(req.message, req.session_id, images=req.images, files=req.files, space=req.space):
                    if event.get("type") == "done":
                        for _fe in flush_session_files(sid):
                            yield f"data: {json.dumps(_fe, ensure_ascii=False)}\n\n"
                    yield f"data: {json.dumps(event, ensure_ascii=False)}\n\n"
            except Exception as e:
                yield f"data: {json.dumps({'type': 'error', 'content': str(e)}, ensure_ascii=False)}\n\n"
            finally:
                if _restore:
                    (
                        agent.provider_name,
                        agent.provider,
                        agent.model,
                        agent._override_temperature,
                        agent.context,
                    ) = _restore
            for _fe in flush_session_files(sid):
                yield f"data: {json.dumps(_fe, ensure_ascii=False)}\n\n"

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


# --- Sessions ---
@app.get("/api/sessions")
async def list_sessions(space: str = "", _=Depends(verify_api_key)):
    """List all sessions with preview text (可按 space 过滤)."""
    result = []
    for s in agent._sessions.values():
        if space and getattr(s, "space", "work") != space:
            continue
        preview = ""
        for m in s.messages:
            if m.get("role") == "user":
                preview = m["content"][:40]
                break
        result.append({
            "id": s.id,
            "preview": preview,
            "message_count": len(s.messages),
            "space": getattr(s, "space", "work"),
        })
    return result


@app.get("/api/sessions/{session_id}")
async def get_session(session_id: str, _=Depends(verify_api_key)):
    """Get full message history for a session."""
    session = agent.get_or_create_session(session_id)
    return {
        "id": session.id,
        "messages": session.messages,
    }


@app.delete("/api/sessions/{session_id}")
async def clear_session(session_id: str, _=Depends(verify_api_key)):
    """Delete a session — requires auth."""
    try:
        agent.clear_session(session_id)
    except Exception as e:
        import logging
        logging.getLogger("lumu").warning("clear_session cleanup failed for %s: %s", session_id, e)
    return {"ok": True}
    return {"ok": True}


@app.post("/api/sessions")
async def create_session(space: str = "work", _=Depends(verify_api_key)):
    """Create a new chat session and return its id (可按 space 指定空间)."""
    new_id = str(uuid.uuid4())
    session = agent.get_or_create_session(new_id, space=space)
    return {"id": session.id, "preview": "", "message_count": 0, "space": session.space}


# 说明：原 /api/auth/login 已移除。
# LUMU 是个人智能体，不设账号体系；访问控制由 core/access_guard 按「暴露面」决定，
# 详见 verify_api_key。该端点此前还内联了开发者凭据默认值，一并清除。


# --- Memory ---
# --- Memory confirmation store（非破坏性、隔离的 JSON 存储，不碰 MemoryManager 表结构）---
_MEM_CONFIRM_PATH = os.path.join(os.getenv("AGENT_HOME", str(Path(__file__).resolve().parent.parent)), "data", "memory_confirmations.json")

def _load_confirms() -> dict:
    try:
        with open(_MEM_CONFIRM_PATH, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {}

def _save_confirm(key: str, confirmed: bool):
    data = _load_confirms()
    if confirmed:
        data[key] = True
    else:
        data.pop(key, None)
    try:
        os.makedirs(os.path.dirname(_MEM_CONFIRM_PATH), exist_ok=True)
        with open(_MEM_CONFIRM_PATH, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    except Exception:
        pass


@app.get("/api/memory")
async def list_memories(category: str = "", space: str = "", _=Depends(verify_api_key)):
    """List all memories, optionally filtered by category and/or space."""
    items = agent.memory.list_all(category if category else None, space if space else None)
    confirms = _load_confirms()
    # 归一化（非破坏性）：确保前端三级置信度分级始终有 importance 字段
    for it in items:
        if not isinstance(it, dict):
            continue
        if it.get("importance") is None:
            md = it.get("metadata") or {}
            it["importance"] = md.get("confidence", 0.6)
        key = it.get("key")
        if key in confirms and confirms[key]:
            it["importance"] = max(float(it.get("importance") or 0), 0.95)
            md = it.get("metadata") or {}
            md["confirmed"] = True
            it["metadata"] = md
            it["confirmed"] = True
    return items


@app.post("/api/memory")
async def save_memory(req: MemoryRequest, _=Depends(verify_api_key)):
    """统一写入路由：store=primary(默认,MemoryManager) 或 semantic(SemanticMemory)。"""
    if req.store == "semantic" and agent.semantic_memory is not None:
        imp = req.importance if req.importance is not None else 0.6
        meta = dict(req.metadata or {})
        if req.confirmed:
            meta["confirmed"] = True
            imp = max(float(imp or 0), 0.95)
        agent.semantic_memory.save(req.key, req.content, req.category, importance=imp, metadata=meta)
    else:
        agent.memory.save(req.key, req.content, req.category)
    if req.confirmed:
        _save_confirm(req.key, True)
    return {"ok": True, "key": req.key, "store": req.store}


@app.delete("/api/memory/{key}")
async def delete_memory(key: str, _=Depends(verify_api_key)):
    """Delete a memory by key."""
    agent.memory.delete(key)
    return {"ok": True, "key": key}


@app.delete("/api/memory/space/{space}")
async def clear_memory_space(space: str, _=Depends(verify_api_key)):
    """清空指定 space 的全部记忆（primary + semantic），用于用户侧隐私重置。"""
    n_primary = 0
    try:
        for it in (agent.memory.list_all(None, space) or []):
            k = it.get("key") if isinstance(it, dict) else None
            if k:
                try:
                    agent.memory.delete(k)
                    n_primary += 1
                except Exception:
                    pass
    except Exception:
        pass
    n_semantic = 0
    try:
        if agent.semantic_memory is not None:
            for it in (agent.semantic_memory.list_all(space=space) or []):
                k = it.get("key") if isinstance(it, dict) else None
                if k:
                    try:
                        agent.semantic_memory.delete(k)
                        n_semantic += 1
                    except Exception:
                        pass
    except Exception:
        pass
    return {"ok": True, "space": space, "primary_deleted": n_primary, "semantic_deleted": n_semantic}


@app.get("/api/memory/search")
async def search_memories(q: str = "", limit: int = 5, _=Depends(verify_api_key)):
    """Search memories by keyword."""
    if not q:
        return []
    return agent.memory.search(q, limit)


@app.get("/api/memory/conflicts")
async def memory_conflicts(_=Depends(verify_api_key)):
    """非破坏性只读：扫描记忆库识别近似重复与可能矛盾的条目。"""
    import re as _re
    items = agent.memory.list_all() or []

    def _tok(s):
        s = (s or "").lower()
        toks = set(_re.findall(r"[a-z0-9]+", s))
        toks.update(_re.findall(r"[一-鿿]", s))
        return toks

    def _jac(a, b):
        if not a or not b:
            return 0.0
        return len(a & b) / len(a | b)

    flags = []
    n = len(items)
    seen = set()
    for i in range(n):
        for j in range(i + 1, n):
            a, b = items[i], items[j]
            jc = _jac(_tok(a.get("content")), _tok(b.get("content")))
            if jc >= 0.55:
                flags.append({
                    "type": "duplicate",
                    "keys": [a.get("key"), b.get("key")],
                    "similarity": round(jc, 2),
                    "detail": "内容高度近似，可能为重复记忆",
                })
                seen.add((i, j))
    # 同分类内的潜在矛盾：含相反极性词
    polarity = [("喜欢", "不喜欢"), ("偏好", "不偏好"), ("要", "不要"),
                ("需要", "不需要"), ("是", "不是"), ("应该", "不应该")]
    for i in range(n):
        for j in range(i + 1, n):
            if (i, j) in seen:
                continue
            a, b = items[i], items[j]
            if a.get("category") != b.get("category"):
                continue
            ca, cb = (a.get("content") or ""), (b.get("content") or "")
            for pos, neg in polarity:
                if (pos in ca and neg in cb) or (neg in ca and pos in cb):
                    flags.append({
                        "type": "conflict",
                        "keys": [a.get("key"), b.get("key")],
                        "detail": f"同分类下出现相反表述（{pos}/{neg}）",
                    })
                    break
    return {"conflicts": flags, "total": n}


# --- 统一记忆总线：聚合多存储层（MemoryManager + SemanticMemory）为单一只读视图 ---
@app.get("/api/memory/unified")
async def memory_unified(space: str = "", _=Depends(verify_api_key)):
    """非破坏性只读：合并 MemoryManager 与 SemanticMemory，按 key 去重并归一化 importance。"""
    confirms = _load_confirms()
    merged = {}
    for it in (agent.memory.list_all(space=space if space else None) or []):
        if not isinstance(it, dict):
            continue
        key = it.get("key")
        if key is None:
            continue
        imp = it.get("importance")
        if imp is None:
            imp = (it.get("metadata") or {}).get("confidence", 0.6)
        if key in confirms and confirms[key]:
            imp = max(float(imp or 0), 0.95)
        merged[key] = {
            "key": key,
            "content": it.get("content", ""),
            "category": it.get("category", "general"),
            "importance": float(imp or 0.6),
            "store": "primary",
            "created_at": it.get("created_at"),
            "confirmed": bool(key in confirms and confirms[key]),
        }
    if agent.semantic_memory is not None:
        for it in (agent.semantic_memory.list_all(limit=2000, space=space if space else None) or []):
            if not isinstance(it, dict):
                continue
            key = it.get("key")
            if key is None:
                continue
            imp = it.get("importance")
            if imp is None:
                imp = (it.get("metadata") or {}).get("confidence", 0.5)
            if key in confirms and confirms[key]:
                imp = max(float(imp or 0), 0.95)
            if key in merged:
                base = merged[key]
                base["importance"] = float(max(base.get("importance", 0), imp or 0))
                base["store"] = "both"
                base["access_count"] = it.get("access_count", 0)
                base["semantic_created_at"] = it.get("created_at")
                if imp is not None:
                    base["confirmed"] = base.get("confirmed") or bool((it.get("metadata") or {}).get("confirmed"))
            else:
                merged[key] = {
                    "key": key,
                    "content": it.get("content", ""),
                    "category": it.get("category", "general"),
                    "importance": float(imp or 0.5),
                    "store": "semantic",
                    "created_at": it.get("created_at"),
                    "access_count": it.get("access_count", 0),
                    "confirmed": bool((it.get("metadata") or {}).get("confirmed")),
                }
    return {"memories": list(merged.values()), "total": len(merged)}


# --- 时间衰减 + 主动遗忘（notes 方案3，非破坏性只读视图 + 显式遗忘动作）---
@app.get("/api/memory/decay")
async def memory_decay(_=Depends(verify_api_key)):
    """非破坏性只读：基于记忆年龄计算衰减后的 importance，标注遗忘候选。"""
    import datetime as _dt
    HALF_LIFE_DAYS = 180.0
    FORGET_THRESHOLD = 0.4
    items = []
    try:
        items = (await memory_unified()).get("memories", [])
    except Exception:
        items = (agent.memory.list_all() or [])
    now = _dt.datetime.now(_dt.timezone.utc)
    out = []
    for m in items:
        created = m.get("created_at") or m.get("semantic_created_at")
        age_days = None
        base_imp = float(m.get("importance") or 0.6)
        if created:
            s = str(created).replace("Z", "").replace("T", " ").strip()
            ct = None
            for fmt in ("%Y-%m-%d %H:%M:%S", "%Y-%m-%d %H:%M:%S.%f", "%Y-%m-%d"):
                try:
                    ct = _dt.datetime.strptime(s.split("+")[0], fmt)
                    break
                except Exception:
                    ct = None
            if ct is not None:
                ct = ct.replace(tzinfo=_dt.timezone.utc)
                age_days = max(0, (now - ct).days)
        decayed = base_imp * (0.5 ** (age_days / HALF_LIFE_DAYS)) if age_days is not None else base_imp
        out.append({
            "key": m.get("key"),
            "importance": round(base_imp, 3),
            "decayed_importance": round(decayed, 3),
            "age_days": age_days,
            "forget_candidate": bool(decayed < FORGET_THRESHOLD),
            "store": m.get("store"),
        })
    forget_count = sum(1 for o in out if o["forget_candidate"])
    return {"decay": out, "total": len(out), "forget_candidates": forget_count,
            "half_life_days": HALF_LIFE_DAYS, "threshold": FORGET_THRESHOLD}


class ForgetRequest(BaseModel):
    keys: list[str] = []
    store: str = "all"  # all / primary / semantic


@app.post("/api/memory/forget")
async def memory_forget(req: ForgetRequest, _=Depends(verify_api_key)):
    """主动遗忘：删除指定 key（默认双库）。显式动作，不自动批量删。"""
    deleted = []
    for key in req.keys:
        if req.store in ("all", "primary") and agent.memory is not None:
            try:
                agent.memory.delete(key)
            except Exception:
                pass
        if req.store in ("all", "semantic") and agent.semantic_memory is not None:
            try:
                agent.semantic_memory.delete(key)
            except Exception:
                pass
        deleted.append(key)
        _save_confirm(key, False)
    return {"ok": True, "deleted": deleted, "store": req.store}

# --- 保守版主动遗忘：归档（可恢复，隔离存储 data/memory_archive.json，非破坏）---
def _archive_path():
    import os as _os
    p = _os.path.join("data", "memory_archive.json")
    _os.makedirs("data", exist_ok=True)
    return p


def _load_archive() -> dict:
    import json as _json, os as _os
    p = _archive_path()
    if not _os.path.exists(p):
        return {}
    try:
        with open(p, encoding="utf-8") as f:
            return _json.load(f)
    except Exception:
        return {}


def _save_archive(d: dict):
    import json as _json
    with open(_archive_path(), "w", encoding="utf-8") as f:
        _json.dump(d, f, ensure_ascii=False, indent=1)


class ArchiveRequest(BaseModel):
    keys: list[str] = []


@app.post("/api/memory/archive")
async def memory_archive(req: ArchiveRequest, _=Depends(verify_api_key)):
    """归档：先完整快照到 data/memory_archive.json，再从双库移除。可随时恢复，显式动作不自动批量。"""
    import datetime as _dt
    uni = await memory_unified()
    by_key = {m["key"]: m for m in uni.get("memories", [])}
    arch = _load_archive()
    archived, missing = [], []
    for key in req.keys:
        m = by_key.get(key)
        if m is None:
            missing.append(key)
            continue
        m = dict(m)
        m["archived_at"] = _dt.datetime.now(_dt.timezone.utc).isoformat()
        arch[key] = m
        if agent.memory is not None:
            try:
                agent.memory.delete(key)
            except Exception:
                pass
        if agent.semantic_memory is not None:
            try:
                agent.semantic_memory.delete(key)
            except Exception:
                pass
        _save_confirm(key, False)
        archived.append(key)
    _save_archive(arch)
    return {"ok": True, "archived": archived, "missing": missing, "total_archived": len(arch)}


@app.get("/api/memory/archived")
async def memory_archived(_=Depends(verify_api_key)):
    """归档区列表（只读）。"""
    arch = _load_archive()
    items = sorted(arch.values(), key=lambda x: x.get("archived_at") or "", reverse=True)
    return {"archived": items, "total": len(items)}


@app.post("/api/memory/restore")
async def memory_restore(req: ArchiveRequest, _=Depends(verify_api_key)):
    """从归档区恢复记忆到原存储，并移出归档。"""
    arch = _load_archive()
    restored, missing = [], []
    for key in req.keys:
        m = arch.get(key)
        if m is None:
            missing.append(key)
            continue
        store = m.get("store") or "primary"
        content = m.get("content", "")
        category = m.get("category", "general")
        imp = float(m.get("importance") or 0.6)
        if store in ("primary", "both") and agent.memory is not None:
            try:
                agent.memory.save(key, content, category)
            except Exception:
                pass
        if store in ("semantic", "both") and agent.semantic_memory is not None:
            try:
                agent.semantic_memory.save(key, content, category, importance=imp, metadata={"restored": True})
            except Exception:
                pass
        if m.get("confirmed"):
            _save_confirm(key, True)
        del arch[key]
        restored.append(key)
    _save_archive(arch)
    return {"ok": True, "restored": restored, "missing": missing, "total_archived": len(arch)}



# --- 记忆智能维护（归纳 / 主动遗忘） ---
@app.post("/api/memory/consolidate")
async def memory_consolidate(_=Depends(verify_api_key)):
    """记忆归纳：把高度相似的记忆合并为一条高层抽象记忆（只新增、不删原）。手动触发。"""
    try:
        n = await agent._memory_consolidate()
        return {"ok": True, "consolidated": n}
    except Exception as e:
        return {"ok": False, "error": str(e)}


@app.post("/api/memory/auto-forget")
async def memory_auto_forget(cap: int = 5, _=Depends(verify_api_key)):
    """主动遗忘（可恢复）：软归档衰减后重要性极低且陈旧的记忆（不硬删，归档到 memory_archive.json）。手动触发。"""
    try:
        n = await agent._auto_forget_memories(cap=max(1, min(cap, 20)))
        return {"ok": True, "archived": n}
    except Exception as e:
        return {"ok": False, "error": str(e)}


# --- Skills ---
class SkillRequest(BaseModel):
    name: str
    description: str
    content: str
    tags: str = ""
    space: str = "work"


@app.get("/api/skills")
async def list_skills(tag: str = "", space: str = "", _=Depends(verify_api_key)):
    """List all saved skills (可按 space 过滤 work/personal)."""
    return agent.skills.list_all(tag if tag else "", space if space else "")


@app.get("/api/skills/{skill_name}")
async def get_skill_api(skill_name: str, _=Depends(verify_api_key)):
    """Get full details of a skill."""
    skill = agent.skills.get(skill_name)
    if not skill:
        raise HTTPException(status_code=404, detail=f"Skill '{skill_name}' not found")
    return skill


@app.post("/api/skills")
async def save_skill_api(req: SkillRequest, _=Depends(verify_api_key)):
    """Save or update a skill."""
    is_new = agent.skills.save(req.name, req.description, req.content, req.tags, req.space)
    return {"ok": True, "name": req.name, "created": is_new}


@app.delete("/api/skills/{skill_name}")
async def delete_skill_api(skill_name: str, _=Depends(verify_api_key)):
    """Delete a skill."""
    if agent.skills.delete(skill_name):
        return {"ok": True, "name": skill_name}
    raise HTTPException(status_code=404, detail=f"Skill '{skill_name}' not found")


@app.get("/api/skills/search")
@app.get("/api/market/skills")
async def market_skills(_=Depends(verify_api_key)):
    import os as _os, shutil as _sh
    base = _os.path.join(_os.path.dirname(_os.path.dirname(_os.path.abspath(__file__))), "skills")
    packs = _os.path.join(base, "packs")
    out = []
    if _os.path.isdir(packs):
        for name in sorted(_os.listdir(packs)):
            pdir = _os.path.join(packs, name)
            if not _os.path.isdir(pdir):
                continue
            meta = {"name": name, "description": "", "installed": _os.path.isdir(_os.path.join(base, name))}
            md = _os.path.join(pdir, "SKILL.md")
            if _os.path.isfile(md):
                try:
                    with io.open(md, encoding="utf-8") as mf:
                        t = mf.read()
                    if t.startswith("---"):
                        fm = t.split("---", 2)
                        if len(fm) >= 3:
                            for ln in fm[1].splitlines():
                                if ln.strip().startswith("description:"):
                                    meta["description"] = ln.split(":", 1)[1].strip()
                except Exception:
                    pass
            out.append(meta)
    return out

@app.post("/api/market/install")
async def market_install(payload: dict, _=Depends(verify_api_key)):
    import os as _os, shutil as _sh
    name = (payload or {}).get("name")
    if not name:
        return {"ok": False, "error": "missing name"}
    base = _os.path.join(_os.path.dirname(_os.path.dirname(_os.path.abspath(__file__))), "skills")
    src = _os.path.join(base, "packs", name)
    dst = _os.path.join(base, name)
    if not _os.path.isdir(src):
        return {"ok": False, "error": "pack not found: " + name}
    if _os.path.exists(dst):
        return {"ok": True, "already": True}
    _sh.copytree(src, dst)
    return {"ok": True, "installed": name}

@app.post("/api/market/publish")
async def market_publish(payload: dict, _=Depends(verify_api_key)):
    import os as _os, re as _re
    p = payload or {}
    name = (p.get("name") or "").strip()
    description = (p.get("description") or "").strip()
    content = (p.get("content") or "").strip()
    triggers = (p.get("triggers") or "").strip()
    if not name or not content:
        return {"ok": False, "error": "name 和 content 不能为空"}
    if not _re.match(r"^[A-Za-z0-9_-]+$", name):
        return {"ok": False, "error": "技能名仅允许字母、数字、中划线、下划线"}
    base = _os.path.join(_os.path.dirname(_os.path.dirname(_os.path.abspath(__file__))), "skills")
    pdir = _os.path.join(base, "packs", name)
    if _os.path.exists(pdir):
        return {"ok": False, "error": "已存在同名技能包"}
    _os.makedirs(pdir, exist_ok=True)
    fm = ["---", "name: " + name, "description: " + description,
          "triggers: " + triggers, "always: false", "---"]
    text = "\n".join(fm) + "\n\n" + content + "\n"
    with io.open(_os.path.join(pdir, "SKILL.md"), "w", encoding="utf-8") as mf:
        mf.write(text)
    return {"ok": True, "published": name}

async def search_skills(q: str = "", limit: int = 5, _=Depends(verify_api_key)):
    """Search skills by keyword."""
    if not q:
        return []
    return agent.skills.search(q, limit)


# --- Cron Jobs ---
@app.get("/api/cron")
async def list_cron_jobs(_=Depends(verify_api_key)):
    """List all scheduled cron jobs."""
    return scheduler.list_jobs()


@app.get("/api/cron/logs")
async def cron_run_logs(job_id: str = "", limit: int = 20, _=Depends(verify_api_key)):
    """Get cron job execution logs."""
    return scheduler.get_run_logs(job_id, limit)


@app.delete("/api/cron/{job_id}")
async def delete_cron_job(job_id: str, _=Depends(verify_api_key)):
    """Delete a cron job."""
    if scheduler.delete_job(job_id):
        return {"ok": True}
    raise HTTPException(status_code=404, detail=f"Job '{job_id}' not found")


# --- HITL 人工审批（模型无权批准，只能人类通过这些端点操作）---
class ApprovalActionRequest(BaseModel):
    feedback: str = ""
    reason: str = ""
    scope: str = "once"  # once | session


@app.get("/api/approvals")
async def list_pending_approvals(_=Depends(verify_api_key)):
    """列出待审批的高风险操作。"""
    from agent.hitl import get_approval_manager
    return {"pending": get_approval_manager().get_pending()}


@app.post("/api/approvals/{action_id}/approve")
async def approve_action(action_id: str, req: ApprovalActionRequest = None, _=Depends(verify_api_key)):
    """人工批准挂起操作并立即执行（唯一合法的批准通道）。"""
    from tools.hitl_tools import approve_and_execute
    from agent.hitl import get_approval_manager
    feedback = req.feedback if req else ""
    scope = (req.scope if req else "once") or "once"
    if scope == "session":
        _mgr = get_approval_manager()
        _act = _mgr.get_status(action_id)
        if _act:
            _mgr.add_session_always_allow(_act.get("session_id"), _act.get("tool_name"))
    result = await approve_and_execute(action_id, feedback=feedback)
    ok = result.startswith("✅") or result.startswith("已批准")
    if not ok:
        raise HTTPException(status_code=404, detail=result)
    return {"ok": True, "result": result}


@app.post("/api/approvals/{action_id}/deny")
async def deny_action(action_id: str, req: ApprovalActionRequest = None, _=Depends(verify_api_key)):
    """人工拒绝挂起操作。"""
    from agent.hitl import get_approval_manager
    reason = req.reason if req else ""
    if get_approval_manager().deny(action_id, reason=reason):
        return {"ok": True}
    raise HTTPException(status_code=404, detail="操作不存在或已过期")


# --- Channels ---
class WebhookMessage(BaseModel):
    text: str
    user_id: str = "webhook_user"
    chat_id: str = "default"
    metadata: dict = {}


@app.post("/api/webhook/{channel_name}")
async def webhook_message(channel_name: str, msg: WebhookMessage, _=Depends(verify_api_key)):
    """Receive a message via webhook channel."""
    try:
        from channels.router import channel_router
        result = await channel_router.route(
            channel_name=channel_name,
            chat_id=msg.chat_id,
            user_id=msg.user_id,
            text=msg.text,
            metadata=msg.metadata,
        )
        return {"ok": True, "reply": result}
    except ImportError:
        raise HTTPException(status_code=501, detail="Channels not available")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# --- Channel callbacks (企业微信 / 飞书 / 钉钉 等回调式渠道) ---
@app.post("/api/channels/{name}/callback")
@limiter.exempt
async def channel_callback(name: str, request: Request):
    """统一回调入口：回调式渠道把平台推送的消息发到这里。

    各适配器自己完成：验签 -> 解析平台消息 -> 调 agent -> 主动回推。
    不依赖我们的 API Key（平台用自身签名校验），故豁免速率限制。
    """
    from channels.router import channel_router
    from core.logging_config import get_logger
    logger = get_logger("api.callback")
    ch = channel_router.get_channel(name)
    if ch is None or not hasattr(ch, "handle_callback"):
        raise HTTPException(status_code=404, detail=f"channel '{name}' not available")
    try:
        raw = await request.body()
        return await ch.handle_callback(raw, request)
    except HTTPException:
        raise
    except Exception as e:
        logger.warning(f"[callback] {name} error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# --- Health ---
@app.get("/api/health")
async def health():
    data = {
        "status": "ok",
        "model": agent.model,
        "provider": agent.provider_name,
        "tools": len(tool_registry.list_tools()),
        "toolsets": list(tool_registry.list_toolsets().keys()),
        "sessions": len(agent._sessions),
        "memories": len(agent.memory.list_all()),
        "skills": len(agent.skills.list_all()),
        "plugins": list(plugin_loader.plugins.keys()),
        "mcp_servers": len(mcp_bridge.get_server_info()),
        "cron_jobs": len(scheduler.list_jobs()),
    }
    try:
        from channels.router import channel_router
        data["channels"] = [c["name"] for c in channel_router.get_status() if c.get("enabled")]
    except Exception:
        data["channels"] = []
    return data


# Mount static files last
app.mount("/static", StaticFiles(directory=str(STATIC_DIR)), name="static")
# 健壮性：确保 assets 目录存在再挂载，避免清前端后整个 app 导入失败导致服务起不来
# 显式 0o755：否则受服务器 umask 影响可能生成 744(其他用户无 x 权限)，
# 导致 lumu 进程无法进入目录读取 JS/CSS -> Starlette 抛 PermissionError -> 静态资源 401 -> 前端白屏
os.makedirs(str(STATIC_DIR / "assets"), mode=0o755, exist_ok=True)
app.mount("/assets", StaticFiles(directory=str(STATIC_DIR / "assets")), name="assets")
# 官方 dashboard 块引用 /avatars/*（nav-user 默认头像）
os.makedirs(str(STATIC_DIR / "avatars"), mode=0o755, exist_ok=True)
app.mount("/avatars", StaticFiles(directory=str(STATIC_DIR / "avatars")), name="avatars")



@app.get("/favicon.svg", include_in_schema=False)
async def favicon_svg():
    """浏览器标签页图标（绿色生命体）。"""
    from fastapi.responses import FileResponse
    p = STATIC_DIR / "favicon.svg"
    if p.exists():
        return FileResponse(str(p), media_type="image/svg+xml")
    raise HTTPException(status_code=404)


@app.get("/favicon.ico", include_in_schema=False)
async def favicon_ico():
    """兼容旧浏览器默认请求 /favicon.ico。"""
    from fastapi.responses import FileResponse
    p = STATIC_DIR / "favicon.svg"
    if p.exists():
        return FileResponse(str(p), media_type="image/svg+xml")
    raise HTTPException(status_code=404)


# --- User Config (Provider API Keys + TTS/STT) ---
from providers.registry import list_providers


class ProviderKeyRequest(BaseModel):
    api_key: str


class TTSSynthesizeRequest(BaseModel):
    text: str
    voice: str | None = None
    provider: str | None = None
    rate: str | None = None
    pitch: str | None = None

class TTSConfigRequest(BaseModel):
    provider: str | None = None
    mimo_api_key: str | None = None


class ModelSwitchRequest(BaseModel):
    provider: str
    model: str


@app.get("/api/config")
async def get_config(_=Depends(verify_api_key)):
    """Get full user configuration (keys masked)."""
    cfg = load_config()
    # Mask API keys for security
    for pname, pcfg in cfg.get("providers", {}).items():
        key = pcfg.get("api_key", "")
        if key and len(key) > 8:
            pcfg["api_key"] = key[:4] + "****" + key[-4:]
    tts = cfg.get("tts", {})
    if tts.get("mimo_api_key") and len(tts["mimo_api_key"]) > 8:
        tts["mimo_api_key"] = tts["mimo_api_key"][:4] + "****"
    # 前端契约补全：旧配置文件可能缺这些键，缺则给默认值，避免设置页读取崩溃
    cfg.setdefault("model_preference", get_model_preference())
    cfg.setdefault("system_prompt", get_system_prompt())
    cfg.setdefault("provider_overrides", {})
    return cfg


def _resolve_enabled_models(provider, configured, has_key):
    """已配置 API Key 的供应商若未显式设置 enabled_models，回退到该供应商的默认模型列表。
    这样前端 TopBar 下拉切换模型时，至少能看到当前供应商可用的全部模型（而非空列表）。
    未配置 key 的供应商不参与 fallback（保持 enabled_models=[]，前端过滤后才显示）。"""
    configured = list(configured) if isinstance(configured, list) else []
    if configured:
        return configured
    if has_key and getattr(provider, "models", None):
        return list(provider.models)
    return []

@app.get("/api/config/providers")
async def get_provider_configs(_=Depends(verify_api_key)):
    """List all providers with their configuration status."""
    providers = list_providers()
    result = []
    for p in providers:
        key = get_provider_key(p.name)
        has_key = bool(key and len(key) > 4)
        result.append({
            "name": p.name,
            "display_name": p.display_name,
            "base_url": p.base_url,
            "models": p.models,
            "context_window": p.context_window,
            "supports_vision": p.supports_vision,
            "api_key_configured": has_key,
            "anthropic_base_url": getattr(p, "anthropic_base_url", "") or "",
            "active_base_url": p.resolve_base_url(),
            "active_protocol": ("anthropic" if ("/anthropic" in p.resolve_base_url() or p.resolve_base_url().rstrip("/").endswith("/api/coding")) else "openai"),
            "api_key_preview": (key[:4] + "****" + key[-4:]) if has_key and len(key) > 8 else ("已配置" if has_key else ""),
            "enabled_models": _resolve_enabled_models(p, get_enabled_models(p.name), has_key),
        })
    return {"providers": result, "total": len(result), "system_prompt": get_system_prompt()}


@app.get("/api/config/providers/{provider_name}/models")
async def detect_provider_models(provider_name: str, _=Depends(verify_api_key)):
    """用已配置密钥调用提供商 /models 接口，自动识别该令牌下真实可用的模型列表。

    识别成功返回 detected=True 与真实模型；失败（无密钥/网络/接口不支持）
    回退到内置静态列表并返回 detected=False。
    """
    from providers.registry import get as get_provider
    p = get_provider(provider_name)
    if not p:
        raise HTTPException(status_code=404, detail=f"Provider '{provider_name}' not found")
    key = get_provider_key(provider_name)
    if not key:
        return {"detected": False, "reason": "未配置 API Key", "models": p.models,
                "enabled_models": get_enabled_models(provider_name)}
    # 用 OpenAI 兼容端点探测（anthropic 端点不提供 /models 列表）
    base = p.resolve_base_url()
    if "/anthropic" in base or base.rstrip("/").endswith("/api/coding"):
        base = p.base_url
    url = base.rstrip("/") + "/models"
    try:
        import httpx
        async with httpx.AsyncClient(timeout=12) as client:
            r = await client.get(url, headers={"Authorization": f"Bearer {key}"})
        if r.status_code != 200:
            return {"detected": False, "reason": f"HTTP {r.status_code}", "models": p.models,
                    "enabled_models": get_enabled_models(provider_name)}
        data = r.json()
        items = data.get("data", data if isinstance(data, list) else [])
        models = []
        for it in items:
            mid = it.get("id") if isinstance(it, dict) else (it if isinstance(it, str) else None)
            if mid and mid not in models:
                models.append(mid)
        if not models:
            return {"detected": False, "reason": "接口未返回模型", "models": p.models,
                    "enabled_models": get_enabled_models(provider_name)}
        return {"detected": True, "models": sorted(models),
                "enabled_models": get_enabled_models(provider_name)}
    except Exception as e:
        return {"detected": False, "reason": str(e)[:120], "models": p.models,
                "enabled_models": get_enabled_models(provider_name)}


class EnabledModelsRequest(BaseModel):
    models: list[str]


@app.post("/api/config/providers/{provider_name}/enabled-models")
async def save_enabled_models(provider_name: str, req: EnabledModelsRequest, _=Depends(verify_api_key)):
    """保存用户勾选启用的模型列表。"""
    from providers.registry import get as get_provider
    if not get_provider(provider_name):
        raise HTTPException(status_code=404, detail=f"Provider '{provider_name}' not found")
    set_enabled_models(provider_name, req.models)
    return {"ok": True, "provider": provider_name, "enabled_models": get_enabled_models(provider_name)}


@app.post("/api/config/provider/{provider_name}/key")
async def set_provider_api_key(provider_name: str, req: ProviderKeyRequest, _=Depends(verify_api_key)):
    """Set API key for a provider."""
    if not req.api_key or len(req.api_key.strip()) < 4:
        raise HTTPException(status_code=400, detail="API key too short")
    set_provider_key(provider_name, req.api_key.strip())
    return {"ok": True, "provider": provider_name, "message": f"{provider_name} API Key 已保存"}


@app.get("/api/config/tts")
async def get_tts_configuration(_=Depends(verify_api_key)):
    """Get TTS configuration."""
    tts = get_tts_config()
    if tts.get("mimo_api_key") and len(tts["mimo_api_key"]) > 8:
        tts["mimo_api_key"] = tts["mimo_api_key"][:4] + "****"
    return tts


@app.post("/api/config/tts")
async def update_tts_configuration(req: TTSConfigRequest, _=Depends(verify_api_key)):
    """Update TTS configuration."""
    set_tts_config(provider=req.provider, mimo_api_key=req.mimo_api_key)
    return {"ok": True, "message": "TTS 配置已更新"}


@app.post("/api/tts/synthesize")
async def tts_synthesize(req: TTSSynthesizeRequest, _=Depends(verify_api_key)):
    """合成文本为语音，返回 base64 音频供前端直接播放（避开 audio 标签无鉴权头难题）。"""
    from tools.tts_stt import handle_tts_synthesize
    import base64, os
    kwargs = {"text": req.text}
    if req.voice:
        kwargs["voice"] = req.voice
    if req.provider:
        kwargs["provider"] = req.provider
    if req.rate:
        kwargs["rate"] = req.rate
    if req.pitch:
        kwargs["pitch"] = req.pitch
    res = handle_tts_synthesize(**kwargs)
    if "error" in res:
        raise HTTPException(status_code=400, detail=res["error"])
    fp = res.get("filepath")
    if not fp or not os.path.exists(fp):
        raise HTTPException(status_code=500, detail="音频文件生成失败")
    with open(fp, "rb") as f:
        b64 = base64.b64encode(f.read()).decode("utf-8")
    return {"ok": True, "audio_base64": b64, "voice": res.get("voice"), "provider": res.get("provider"), "format": "mp3"}


@app.get("/tts-demo")
async def tts_demo_page():
    from fastapi.responses import FileResponse
    base = os.getenv("AGENT_HOME", str(Path(__file__).resolve().parent.parent))
    return FileResponse(os.path.join(base, "api", "tts_demo.html"), media_type="text/html")


# --- 语音输入（STT）---
@app.post("/api/stt/transcribe")
async def stt_transcribe_api(
    audio: UploadFile = File(...),
    language: str = Form("zh"),
    _=Depends(verify_api_key),
):
    """浏览器录音 → 文本。本地 faster-whisper 推理，音频不出服务器。"""
    import os as _os, tempfile as _tf
    from starlette.concurrency import run_in_threadpool
    data = await audio.read()
    if not data:
        raise HTTPException(status_code=400, detail="空音频")
    if len(data) > 25 * 1024 * 1024:
        raise HTTPException(status_code=413, detail="音频过大（上限 25MB）")
    ext = _os.path.splitext(audio.filename or "")[1] or ".webm"
    fd, tmp = _tf.mkstemp(suffix=ext); _os.close(fd)
    with open(tmp, "wb") as fh:
        fh.write(data)
    try:
        from tools import stt_local
        if not stt_local.is_available():
            raise HTTPException(status_code=503, detail="语音识别引擎未安装（faster-whisper）")
        res = await run_in_threadpool(stt_local.transcribe, tmp, language)
        if not res.get("ok"):
            raise HTTPException(status_code=500, detail=res.get("error", "识别失败"))
        return res
    finally:
        try:
            _os.remove(tmp)
        except Exception:
            pass


@app.get("/api/stt/status")
async def stt_status_api():
    """语音识别可用性探测，供前端决定是否显示麦克风按钮。"""
    from tools import stt_local
    return {"available": stt_local.is_available(), "model": stt_local.MODEL_SIZE,
            "provider": "local_faster_whisper"}


@app.get("/api/config/stt")
async def get_stt_configuration(_=Depends(verify_api_key)):
    """Get STT configuration."""
    return get_stt_config()


@app.post("/api/config/model")
async def switch_model(req: ModelSwitchRequest, _=Depends(verify_api_key)):
    """Switch the active model (and optionally provider)."""
    from providers.registry import get as get_provider

    provider = get_provider(req.provider)
    if not provider:
        raise HTTPException(status_code=404, detail=f"Provider '{req.provider}' not found")
    allowed = set(provider.models) | set(get_enabled_models(req.provider))
    if req.model not in allowed:
        raise HTTPException(status_code=400, detail=f"Model '{req.model}' not available for '{req.provider}'")

    # Update agent
    if agent.provider_name != req.provider:
        agent.provider_name = req.provider
        agent.provider = provider
        # Rebuild context engine for new provider's context window
        from agent.context import ContextEngine
        agent.context = ContextEngine(context_window=provider.context_window)
    agent.model = req.model

    # Persist model preference
    set_model_preference(req.provider, req.model)

    return {
        "ok": True,
        "provider": req.provider,
        "model": req.model,
        "message": f"已切换到 {provider.display_name} / {req.model}"
    }


class SystemPromptRequest(BaseModel):
    system_prompt: str


@app.post("/api/config/system-prompt")
async def set_system_prompt_api(req: SystemPromptRequest, _=Depends(verify_api_key)):
    """Save custom system prompt."""
    try:
        set_system_prompt(req.system_prompt)
        return {"status": "ok"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# --- Embedding Config ---
class EmbeddingConfigRequest(BaseModel):
    base_url: str = ""
    model: str = ""
    api_key: str = ""


@app.get("/api/config/embedding")
async def get_embedding_config_api(_=Depends(verify_api_key)):
    """Get embedding API configuration."""
    cfg = get_embedding_config()
    # Mask API key for security
    if cfg.get("api_key") and len(cfg["api_key"]) > 8:
        cfg["api_key_masked"] = cfg["api_key"][:4] + "****" + cfg["api_key"][-4:]
        cfg["api_key"] = True
    else:
        cfg["api_key_masked"] = ""
        cfg["api_key"] = bool(cfg.get("api_key"))
    return cfg


@app.post("/api/config/embedding")
async def set_embedding_config_api(req: EmbeddingConfigRequest, _=Depends(verify_api_key)):
    """Save embedding API configuration."""
    try:
        # Don't overwrite existing key with empty string
        existing = get_embedding_config()
        api_key = req.api_key if req.api_key.strip() else existing.get("api_key", "")
        set_embedding_config(
            api_key=api_key,
            base_url=req.base_url.strip(),
            model=req.model.strip(),
        )
        cfg = get_embedding_config()
        if cfg.get("api_key") and len(cfg["api_key"]) > 8:
            cfg["api_key_masked"] = cfg["api_key"][:4] + "****" + cfg["api_key"][-4:]
            cfg["api_key"] = True
        else:
            cfg["api_key_masked"] = ""
            cfg["api_key"] = bool(cfg.get("api_key"))
        return cfg
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/kb/reembed")
async def reembed_knowledge(space: str = "", _=Depends(verify_api_key)):
    """Re-embed all knowledge entries using current embedding method (按 space 隔离)."""
    try:
        from knowledge.base import KnowledgeBase
        kb_path = _kb_db_path(space)
        kb = KnowledgeBase(db_path=kb_path)
        result = kb.reembed_all()
        from knowledge.embedding import get_embedding_info
        result["embedding_info"] = get_embedding_info()
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


def _kb_db_path(space: str = ""):
    """按空间隔离知识库：knowledge_{space}.db，默认 work。"""
    sp = space or "work"
    home = os.getenv("AGENT_HOME", str(Path(__file__).resolve().parent.parent))
    return os.path.join(home, "data", f"knowledge_{sp}.db")


# --- Knowledge Base Documents ---

@app.get("/api/kb/documents")
async def list_kb_documents(space: str = "", _=Depends(verify_api_key)):
    """List knowledge base documents with stats (按 space 隔离)."""
    try:
        from knowledge.base import KnowledgeBase
        kb_path = _kb_db_path(space)
        kb = KnowledgeBase(db_path=kb_path)
        entries = kb.list_entries()
        stats = kb.stats()
        return {"documents": entries, "stats": stats}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


def _extract_doc_text(raw: bytes, ext: str, filename: str):
    """抽取文档纯文本（真实解析，非占位）。

    返回 (text, page_marks)：
      - text: 全文（PDF 按页拼接，带换行）
      - page_marks: [(char_offset, page_no), ...]，标记某段文本起始于第几页（PDF 用，其余默认第 1 页）
    解析失败时返回 ("", [(0, 1)])，调用方据此报空文档而非崩溃。
    """
    if ext in (".txt", ".md"):
        text = raw.decode("utf-8", errors="ignore")
        return text, [(0, 1)]

    if ext == ".pdf":
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(raw))
            parts = []
            marks = []
            for i, page in enumerate(reader.pages):
                t = page.extract_text() or ""
                if t.strip():
                    marks.append((len("".join(parts)), i + 1))
                    parts.append(t + "\n")
            return "".join(parts), marks
        except Exception as e:
            print(f"[KB] PDF 解析失败 {filename}: {e}")
            return "", [(0, 1)]

    if ext == ".docx":
        try:
            from docx import Document
            doc = Document(io.BytesIO(raw))
            paras = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
            # 同时抽取表格文本，避免表格内容被漏掉
            for tbl in doc.tables:
                for row in tbl.rows:
                    cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                    if cells:
                        paras.append(" | ".join(cells))
            return "\n".join(paras), [(0, 1)]
        except Exception as e:
            print(f"[KB] DOCX 解析失败 {filename}: {e}")
            return "", [(0, 1)]

    # 兜底：当作 utf-8 文本
    return raw.decode("utf-8", errors="ignore"), [(0, 1)]


def _chunk_text(text: str, size: int = 1200, overlap: int = 200):
    """把长文按「尽量在段落/换行处断开」切成若干块。

    返回 [{"text": str, "start": int}, ...]，start 为块在原文中的字符偏移（用于页码回溯）。
    单段落超长时会硬切，避免一条 entry 过长。
    """
    text = (text or "").strip()
    if not text:
        return []
    chunks = []
    i = 0
    n = len(text)
    while i < n:
        end = min(i + size, n)
        if end < n:
            # 在窗口内靠后位置找换行作为断点，尽量不劈开句子
            nl = text.find("\n", i + int(size * 0.55), end)
            if nl != -1:
                end = nl
        seg = text[i:end].strip()
        if seg:
            chunks.append({"text": seg, "start": i})
        if end >= n:
            break
        i = max(end - overlap, i + 1)
    return chunks


def _page_for_chunk(page_marks, start):
    """根据字符偏移回溯该块所属页码（仅 PDF 有意义）。"""
    if not page_marks:
        return None
    page = page_marks[0][1]
    for off, pg in page_marks:
        if off <= start:
            page = pg
        else:
            break
    return page


@app.post("/api/kb/documents")
async def upload_kb_document(file: UploadFile = File(...), space: str = "", _=Depends(verify_api_key)):
    """Upload a document to the knowledge base (按 space 隔离).

    真实解析 txt/md/pdf/docx，按语义与长度拆分成多条知识 entry 入库，
    返回解析统计：识别字符数、拆分段数、入库条数。
    """
    try:
        from knowledge.base import KnowledgeBase
        kb_path = _kb_db_path(space)
        kb = KnowledgeBase(db_path=kb_path)

        raw = await file.read()
        filename = file.filename or "untitled"
        ext = os.path.splitext(filename)[1].lower()
        base_title = os.path.splitext(filename)[0]

        # 1) 抽取文本
        text, page_marks = _extract_doc_text(raw, ext, filename)

        # 2) 分块
        chunks = _chunk_text(text, size=1200, overlap=200) if text.strip() else []

        # 3) 逐块入库
        added = 0
        for i, ch in enumerate(chunks):
            title = base_title if len(chunks) == 1 else f"{base_title}（第{i+1}/{len(chunks)}段）"
            meta = {
                "source_file": filename,
                "chunk_index": i,
                "chunk_total": len(chunks),
            }
            page_no = _page_for_chunk(page_marks, ch["start"])
            if page_no is not None:
                meta["page"] = page_no
            kb.add(
                title=title,
                content=ch["text"],
                category="uploaded",
                tags=["kb-doc", base_title],
                source=filename,
                metadata=meta,
            )
            added += 1

        if added == 0:
            # 解析为空（扫描件 PDF / 空文件）：明确告知，避免用户误以为入库成功
            return {
                "status": "empty",
                "filename": filename,
                "chars": 0,
                "chunks": 0,
                "entries": 0,
                "message": "文档未解析出可读文本（可能是扫描件图片 PDF 或空文件）",
            }

        return {
            "status": "ok",
            "filename": filename,
            "chars": len(text),
            "chunks": len(chunks),
            "entries": added,
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/api/kb/documents/{entry_id}")
async def delete_kb_document(entry_id: str, space: str = "", _=Depends(verify_api_key)):
    """Delete a knowledge base document by entry id (按 space 隔离)."""
    try:
        from knowledge.base import KnowledgeBase
        kb_path = _kb_db_path(space)
        kb = KnowledgeBase(db_path=kb_path)
        result = kb.delete(entry_id)
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/kb/stats")
async def kb_stats(space: str = "", _=Depends(verify_api_key)):
    """Get knowledge base statistics and categories (按 space 隔离)."""
    try:
        from knowledge.base import KnowledgeBase
        kb_path = _kb_db_path(space)
        kb = KnowledgeBase(db_path=kb_path)
        stats = kb.stats()
        categories = kb.list_categories()
        return {"stats": stats, "categories": categories}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# --- Inference Parameters ---

class ParamsRequest(BaseModel):
    temperature: float = 0.7
    top_p: float = 0.9


@app.get("/api/config/params")
async def get_inference_params(_=Depends(verify_api_key)):
    """Get inference parameters (temperature, top_p)."""
    try:
        cfg = load_config()
        params = cfg.get("inference_params", {})
        return {
            "temperature": params.get("temperature", 0.7),
            "top_p": params.get("top_p", 0.9),
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/config/params")
async def save_inference_params(req: ParamsRequest, _=Depends(verify_api_key)):
    """Save inference parameters (temperature, top_p) to user config."""
    try:
        cfg = load_config()
        cfg.setdefault("inference_params", {})
        cfg["inference_params"]["temperature"] = req.temperature
        cfg["inference_params"]["top_p"] = req.top_p
        save_config(cfg)
        return {"status": "ok", "temperature": req.temperature, "top_p": req.top_p}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# --- System Prompt (GET) ---

@app.get("/api/config/system-prompt")
async def get_system_prompt_api(_=Depends(verify_api_key)):
    """Get current system prompt."""
    try:
        return {"system_prompt": get_system_prompt()}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# --- Knowledge Graph (sphere) ---

@app.get("/api/kb/graph")
async def kb_graph(space: str = "", _=Depends(verify_api_key)):
    """返回知识记忆生命球体的节点/边/统计，供前端 Canvas2D 渲染（knowledge 按 space 隔离）。"""
    import os, sqlite3, json, math
    from collections import defaultdict, Counter
    base = os.getenv("AGENT_HOME", str(Path(__file__).resolve().parent.parent))
    data_dir = os.path.join(base, "data")
    kb_file = f"knowledge_{space or 'work'}.db"

    def _vec(blob):
        if blob is None:
            return None
        if isinstance(blob, (bytes, bytearray)):
            try:
                import array
                a = array.array("f")
                a.frombytes(bytes(blob))
                if len(a) > 0:
                    return list(a)
            except Exception:
                pass
            return None
        try:
            arr = json.loads(blob) if isinstance(blob, str) else blob
            if isinstance(arr, list) and len(arr) > 0:
                return [float(x) for x in arr]
        except Exception:
            pass
        return None

    def _conn(name):
        return sqlite3.connect(os.path.join(data_dir, name))

    nodes = []
    vecs = {}
    try:
        c = _conn(kb_file)
        for r in c.execute("SELECT id, title, content, category, embedding FROM knowledge"):
            v = _vec(r[4]); nid = "k:%s" % r[0]
            nodes.append({"id": nid, "type": "knowledge", "label": r[1] or "知识", "text": (r[2] or "")[:240], "category": r[3] or ""})
            if v: vecs[nid] = v
        c.close()
    except Exception as e:
        logger_auth.warning("kb_graph_knowledge", err=str(e))
    try:
        c = _conn("memory.db")
        for r in c.execute("SELECT id, key, content, category, graph_vec FROM memories"):
            v = _vec(r[4]); nid = "m:%s" % r[0]
            nodes.append({"id": nid, "type": "user_memory", "label": r[1] or (r[2] or "")[:30], "text": (r[2] or "")[:240], "category": r[3] or ""})
            if v: vecs[nid] = v
        c.close()
    except Exception as e:
        logger_auth.warning("kb_graph_memory", err=str(e))
    try:
        c = _conn("semantic_memory.db")
        for r in c.execute("SELECT id, key, content, category, graph_vec FROM semantic_memories"):
            v = _vec(r[4]); nid = "s:%s" % r[0]
            nodes.append({"id": nid, "type": "semantic", "label": r[1] or (r[2] or "")[:30], "text": (r[2] or "")[:240], "category": r[3] or ""})
            if v: vecs[nid] = v
        for r in c.execute("SELECT id, event_type, description, details, graph_vec FROM episodic_events"):
            v = _vec(r[4]); nid = "e:%s" % r[0]
            nodes.append({"id": nid, "type": "episodic", "label": r[1] or "事件", "text": ((r[2] or "") + " " + (r[3] or ""))[:240], "category": ""})
            if v: vecs[nid] = v
        c.close()
    except Exception as e:
        logger_auth.warning("kb_graph_semantic", err=str(e))
    try:
        c = _conn("lessons.db")
        for r in c.execute("SELECT id, title, lesson_type, description, graph_vec FROM lessons"):
            v = _vec(r[4]); nid = "l:%s" % r[0]
            nodes.append({"id": nid, "type": "lesson", "label": r[1] or "经验", "text": (r[3] or "")[:240], "category": r[2] or ""})
            if v: vecs[nid] = v
        c.close()
    except Exception as e:
        logger_auth.warning("kb_graph_lessons", err=str(e))

    # 斐波那契球面坐标，保证前端 3D 投影有有效 x/y/z
    N = len(nodes)
    golden = math.pi * (3 - math.sqrt(5))
    for i, n in enumerate(nodes):
        if N == 1:
            n["x"], n["y"], n["z"] = 0.0, 0.0, 1.0
        else:
            yy = 1 - (i / (N - 1)) * 2
            rad = math.sqrt(max(0.0, 1 - yy * yy))
            theta = golden * i
            n["x"] = round(math.cos(theta) * rad, 4)
            n["y"] = round(yy, 4)
            n["z"] = round(math.sin(theta) * rad, 4)

    node_by_id = {n["id"]: n for n in nodes}
    type_of = {n["id"]: n["type"] for n in nodes}

    edges = []
    CROSS_THRESH = 0.42
    try:
        import numpy as np
        ids = list(vecs.keys())
        if ids:
            M = np.array([vecs[i] for i in ids], dtype=np.float32)
            norms = np.linalg.norm(M, axis=1, keepdims=True)
            Mn = M / (norms + 1e-9)
            S = (Mn @ Mn.T).tolist()
            intra_cap = defaultdict(list)
            per_node_cross = defaultdict(list)
            N = len(ids)
            for i in range(N):
                for j in range(i + 1, N):
                    s = S[i][j]
                    if s < 0.3:
                        continue
                    ti = type_of[ids[i]]; tj = type_of[ids[j]]
                    if ti == tj:
                        if s > 0.3:
                            intra_cap[ids[i]].append((s, ids[j]))
                            intra_cap[ids[j]].append((s, ids[i]))
                    elif s >= CROSS_THRESH:
                        per_node_cross[ids[i]].append((s, ids[j]))
                        per_node_cross[ids[j]].append((s, ids[i]))
            for nid, lst in intra_cap.items():
                lst.sort(reverse=True)
                for s, other in lst[:3]:
                    edges.append({"source": nid, "target": other, "rel": "related", "weight": round(s, 3)})
            for nid, lst in per_node_cross.items():
                lst.sort(reverse=True)
                for s, other in lst[:4]:
                    edges.append({"source": nid, "target": other, "rel": "semantic", "weight": round(s, 3)})
    except Exception as e:
        logger_auth.warning("kb_graph_sim", err=str(e))

    # 手动语义关系 (P1-3)
    manual_pairs = set()
    try:
        c = _conn("graph.db")
        c.execute("CREATE TABLE IF NOT EXISTS graph_relations(id INTEGER PRIMARY KEY AUTOINCREMENT, source TEXT, target TEXT, rel TEXT, note TEXT, created_at TEXT)")
        for r in c.execute("SELECT id, source, target, rel, note FROM graph_relations"):
            rid, src, tgt, rel, note = r
            if src in node_by_id and tgt in node_by_id:
                edges.append({"source": src, "target": tgt, "rel": rel, "manual": True, "rid": rid, "note": note})
                manual_pairs.add((src, tgt) if src < tgt else (tgt, src))
        c.close()
    except Exception as e:
        logger_auth.warning("kb_graph_manual", err=str(e))

    # 手动边优先：剔除同 pair 的自动边
    if manual_pairs:
        edges = [e for e in edges if e.get("manual") or ((e["source"], e["target"]) if e["source"] < e["target"] else (e["target"], e["source"])) not in manual_pairs]

    # 节点重要度：类型基线 + 连接度归一化（驱动前端节点点半径）
    deg = Counter()
    for e in edges:
        deg[e["source"]] = deg.get(e["source"], 0) + 1
        deg[e["target"]] = deg.get(e["target"], 0) + 1
    type_base = {"knowledge": 0.8, "user_memory": 0.65, "semantic": 0.5, "episodic": 0.4, "lesson": 0.6}
    mdr = math.sqrt(max(deg.values())) if deg else 1.0
    for n in nodes:
        d = deg.get(n["id"], 0)
        norm = (math.sqrt(d) / mdr) if mdr > 0 else 0.0
        n["importance"] = round(min(1.2, type_base.get(n["type"], 0.5) + norm * 0.5), 3)

    stats = {
        "total": len(nodes),
        "edges": len(edges),
        "knowledge": sum(1 for n in nodes if n["type"] == "knowledge"),
        "user_memory": sum(1 for n in nodes if n["type"] == "user_memory"),
        "semantic": sum(1 for n in nodes if n["type"] == "semantic"),
        "episodic": sum(1 for n in nodes if n["type"] == "episodic"),
        "lesson": sum(1 for n in nodes if n["type"] == "lesson"),
    }
    return {"nodes": nodes, "edges": edges, "stats": stats}


class RelationRequest(BaseModel):
    source: str
    target: str
    rel: str
    note: str | None = None


@app.post("/api/kb/graph/relations")
async def add_relation(req: RelationRequest, _=Depends(verify_api_key)):
    """新增手动语义关系 (P1-3)。"""
    import os, sqlite3
    from datetime import datetime
    base = os.getenv("AGENT_HOME", str(Path(__file__).resolve().parent.parent))
    c = sqlite3.connect(os.path.join(base, "data", "graph.db"))
    c.execute("CREATE TABLE IF NOT EXISTS graph_relations(id INTEGER PRIMARY KEY AUTOINCREMENT, source TEXT, target TEXT, rel TEXT, note TEXT, created_at TEXT)")
    cur = c.execute("INSERT INTO graph_relations(source, target, rel, note, created_at) VALUES(?,?,?,?,?)",
                    (req.source, req.target, req.rel, req.note, datetime.utcnow().isoformat()))
    rid = cur.lastrowid
    c.commit(); c.close()
    return {"id": rid, "source": req.source, "target": req.target, "rel": req.rel, "manual": True}


@app.delete("/api/kb/graph/relations/{rel_id}")
async def del_relation(rel_id: int, _=Depends(verify_api_key)):
    """删除手动语义关系 (P1-3)。"""
    import os, sqlite3
    base = os.getenv("AGENT_HOME", str(Path(__file__).resolve().parent.parent))
    c = sqlite3.connect(os.path.join(base, "data", "graph.db"))
    c.execute("DELETE FROM graph_relations WHERE id=?", (rel_id,))
    c.commit(); c.close()
    return {"ok": True}


@app.get("/api/timeline")
async def get_timeline(limit: int = 300, offset: int = 0, _=Depends(verify_api_key)):
    """会话事件时间线 (P1-4)：把每轮展开为 user/tool/assistant 事件。"""
    import os, sqlite3, json
    base = os.getenv("AGENT_HOME", str(Path(__file__).resolve().parent.parent))
    c = sqlite3.connect(os.path.join(base, "data", "interactions.db"))
    try:
        total = list(c.execute("SELECT COUNT(*) FROM interactions"))[0][0]
    except Exception:
        total = 0
    rows = list(c.execute(
        "SELECT id, timestamp, user_msg, assistant_msg, tools_used, outcome, score FROM interactions ORDER BY id DESC LIMIT ? OFFSET ?",
        (limit, offset)))
    events = []
    for r in rows:
        iid, ts, um, am, tu, outcome, score = r
        events.append({"id": "%s-u" % iid, "turn_id": iid, "kind": "user", "ts": ts, "text": um or "", "meta": {}})
        try:
            tools = json.loads(tu) if tu else []
        except Exception:
            tools = []
        if isinstance(tools, list):
            for t in tools:
                if isinstance(t, str):
                    tname = t
                elif isinstance(t, dict):
                    tname = t.get("name") or t.get("tool") or "工具"
                else:
                    tname = str(t)
                events.append({"id": "%s-t" % iid, "turn_id": iid, "kind": "tool", "ts": ts, "text": tname, "meta": {}})
        events.append({"id": "%s-a" % iid, "turn_id": iid, "kind": "assistant", "ts": ts, "text": am or "", "meta": {"outcome": outcome, "score": score}})
    c.close()
    return {"events": events, "total": total, "limit": limit, "offset": offset}


# --- Learning loop (SelfLearningEngine) ---
@app.get("/api/learning/insights")
async def learning_insights(_=Depends(verify_api_key)):
    """只读：返回失败模式、统计与近期经验教训（持续学习闭环的可见层）。"""
    try:
        from agent.learner import LearningEngine
        eng = LearningEngine()
        stats = eng.tracker.get_stats()
        failures = eng.tracker.get_failure_patterns(limit=10)
        lessons = eng.lessons_db.get_all(limit=20)
        return {"stats": stats, "failure_patterns": failures, "lessons": lessons}
    except Exception as e:
        return {"stats": {}, "failure_patterns": [], "lessons": [], "error": str(e)}


@app.post("/api/learning/scan")
async def learning_scan(_=Depends(verify_api_key)):
    """主动扫描近期会话，喂入 SelfLearningEngine 形成持续学习闭环。"""
    try:
        from agent.learner import LearningEngine
        eng = LearningEngine()
        sessions = list(agent._sessions.values())[-20:] if hasattr(agent, "_sessions") else []
        count = 0
        for s in sessions:
            msgs = getattr(s, "messages", None) or []
            user_msgs = [m for m in msgs if getattr(m, "role", "") == "user"]
            asst_msgs = [m for m in msgs if getattr(m, "role", "") == "assistant"]
            if not user_msgs or not asst_msgs:
                continue
            user_msg = user_msgs[0].content if hasattr(user_msgs[0], "content") else ""
            asst_msg = asst_msgs[-1].content if hasattr(asst_msgs[-1], "content") else ""
            tool_calls = []
            for m in msgs:
                for tc in (getattr(m, "tool_calls", None) or []):
                    name = getattr(tc, "name", None) or (
                        getattr(tc, "function", None) and getattr(tc.function, "name"))
                    err = getattr(tc, "error", None)
                    if name:
                        tool_calls.append({"name": name, "error": err})
            # 仅当有真实工具调用或明显成功时才记录
            eng.analyze_interaction(user_msg, asst_msg, tool_calls=tool_calls,
                                    outcome="success" if tool_calls else "partial")
            count += 1
        return {"ok": True, "scanned": count}
    except Exception as e:
        return {"ok": False, "error": str(e)}


# --- Skills auto-extraction ---
class SessionExtractRequest(BaseModel):
    session_id: str = ""


async def _llm_chat(prompt: str, max_tokens: int = 2000) -> str:
    """复用 reasoning.py 的 LLM 调用模式（OpenAI 兼容）。"""
    from openai import AsyncOpenAI
    from providers.registry import get as get_provider
    import config
    provider = get_provider(config.DEFAULT_PROVIDER)
    api_key = os.getenv(provider.api_key_env, "")
    client = AsyncOpenAI(api_key=api_key, base_url=provider.base_url)
    resp = await client.chat.completions.create(
        model=config.DEFAULT_MODEL,
        messages=[{"role": "user", "content": prompt}],
        max_tokens=max_tokens, temperature=0.4,
    )
    return resp.choices[0].message.content or ""


@app.post("/api/skills/extract-from-session")
async def extract_skill_from_session(req: SessionExtractRequest, _=Depends(verify_api_key)):
    """从一次会话的工具编排链自动起草可复用技能（直接保存为技能草稿）。"""
    try:
        if not req.session_id:
            return {"ok": False, "error": "session_id required"}
        session = agent.get_or_create_session(req.session_id)
        msgs = getattr(session, "messages", []) or []
        tool_chain = []
        user_prompt = ""
        for m in msgs:
            role = getattr(m, "role", "")
            if role == "user" and not user_prompt:
                c = getattr(m, "content", "")
                user_prompt = c if isinstance(c, str) else ""
            if role == "assistant":
                for tc in (getattr(m, "tool_calls", None) or []):
                    fn = getattr(tc, "function", None)
                    name = getattr(tc, "name", None) or (fn and getattr(fn, "name"))
                    args = fn and getattr(fn, "arguments", "")
                    if name:
                        tool_chain.append({"name": name, "args": args})
        if not tool_chain:
            return {"ok": False, "error": "该会话未检测到工具调用，无法提取技能"}
        chain_text = "\n".join(
            f"{i+1}. {t['name']}" + (f"({t['args']})" if t.get("args") else "")
            for i, t in enumerate(tool_chain))
        prompt = (
            "你是技能提取器。根据下面的对话意图与工具编排链，起草一个可复用技能。\n"
            f"用户意图：{user_prompt[:400]}\n"
            f"工具编排链：\n{chain_text}\n\n"
            "只输出 JSON：{\"name\":\"snake_case_name\",\"description\":\"何时使用\","
            "\"content\":\"逐步操作 markdown\",\"tags\":\"逗号分隔标签\"}"
        )
        text = await _llm_chat(prompt, max_tokens=1500)
        import re as _re
        m = _re.search(r"\{[\s\S]*\}", text)
        data = json.loads(m.group()) if m else {}
        name = data.get("name")
        if not name:
            return {"ok": False, "error": "LLM 未返回有效技能"}
        _sp = getattr(session, "space", "work") or "work"
        _tags = data.get("tags", "")
        if _sp not in _tags:
            _tags = (_tags + "," + _sp).strip(",")
        is_new = agent.skills.save(name, data.get("description", ""),
                                   data.get("content", ""), _tags, _sp)
        return {"ok": True, "name": name, "created": is_new, "space": _sp, "skill": data}
    except Exception as e:
        return {"ok": False, "error": str(e)}


# --- Feedback ---

class FeedbackRequest(BaseModel):
    session_id: str = ""
    message_index: int = -1
    feedback: str  # "like" or "dislike"（取消时传空字符串 ""）
    message: str = ""   # 被评价的回复内容（用于记忆摘要）
    prompt: str = ""    # 对应的用户问题（用于记忆摘要）
    space: str = ""     # 空间，便于后续按空间隔离分析


@app.post("/api/feedback")
async def submit_feedback(req: FeedbackRequest, _=Depends(verify_api_key)):
    """记录消息赞/踩，并落到记忆系统（赞=认可 / 踩=不认可），供智能体后续学习进化。"""
    try:
        logger = get_logger("feedback")
        logger.info(
            "message_feedback",
            session_id=req.session_id,
            message_index=req.message_index,
            feedback=req.feedback,
        )
        # 落地到语义记忆：让智能体在相关话题时能召回“用户的认可/不认可”，从而自我修正
        if req.feedback in ("like", "dislike") and agent is not None:
            try:
                sm = getattr(agent, "semantic_memory", None)
                if sm is not None:
                    verb = "认可" if req.feedback == "like" else "不认可"
                    content = "用户%s的回答" % verb
                    if req.prompt:
                        content += "（针对问题：%s）" % req.prompt[:200]
                    if req.message:
                        content += "——回答内容：%s" % req.message[:300]
                    key = "feedback_%s_%d" % (req.feedback, int(time.time() * 1000))
                    meta = {
                        "type": "feedback",
                        "rating": req.feedback,
                        "session_id": req.session_id,
                        "message_index": req.message_index,
                        "space": req.space,
                        "prompt": req.prompt[:200],
                        "message": req.message[:500],
                        "ts": time.time(),
                    }
                    sm.save(key, content, "feedback", importance=0.7, metadata=meta)
            except Exception as _fm_e:
                logger.warning("feedback memory save failed: %s" % _fm_e)
        return {"status": "ok", "message": "Feedback recorded"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# --- 结构化产品反馈（个人市场反馈闭环，供 B 端决策）---
import os as _fb_os
import sqlite3 as _fb_sqlite


def _feedback_db():
    _base = _fb_os.path.dirname(_fb_os.path.dirname(_fb_os.path.abspath(__file__)))
    _db = _fb_os.path.join(_base, "data", "feedback.db")
    _c = _fb_sqlite.connect(_db)
    _c.execute(
        """CREATE TABLE IF NOT EXISTS feedback (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            ts TEXT DEFAULT (datetime('now')),
            category TEXT DEFAULT 'suggest',   -- suggest | bug | praise | question
            rating INTEGER DEFAULT 0,          -- 1-5 星
            content TEXT NOT NULL,
            feature TEXT DEFAULT '',
            page TEXT DEFAULT '',
            contact TEXT DEFAULT '',
            user_agent TEXT DEFAULT '',
            status TEXT DEFAULT 'new'          -- new | reviewed | resolved
        )"""
    )
    _c.commit()
    return _c


class ProductFeedbackRequest(BaseModel):
    category: str = "suggest"   # suggest | bug | praise | question
    rating: int = 0             # 1-5 星
    content: str                # 必填
    feature: str = ""
    page: str = ""
    contact: str = ""


@app.post("/api/feedback/submit")
async def submit_product_feedback(req: ProductFeedbackRequest, request: Request):
    """公开的产品级反馈入口（个人用户无需密钥）。落结构化库，供后续分析做 B 端决策。"""
    content = (req.content or "").strip()
    if not content:
        raise HTTPException(status_code=400, detail="content required")
    if req.rating < 0 or req.rating > 5:
        raise HTTPException(status_code=400, detail="rating must be 0-5")
    try:
        c = _feedback_db()
        c.execute(
            "INSERT INTO feedback (category, rating, content, feature, page, contact, user_agent) VALUES (?,?,?,?,?,?,?)",
            (req.category, req.rating, content[:2000], req.feature[:80], req.page[:120],
             req.contact[:120], (request.headers.get("user-agent") or "")[:200]),
        )
        c.commit(); c.close()
        return {"status": "ok", "message": "感谢反馈，我们已收到"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/feedback/admin")
async def list_feedback(category: str = "", status: str = "", limit: int = 100, _=Depends(verify_api_key)):
    """后台查看已收集反馈（受 API key 保护）。"""
    try:
        c = _feedback_db()
        sql = "SELECT id, ts, category, rating, content, feature, page, contact, status FROM feedback"
        clauses, params = [], []
        if category:
            clauses.append("category=?"); params.append(category)
        if status:
            clauses.append("status=?"); params.append(status)
        if clauses:
            sql += " WHERE " + " AND ".join(clauses)
        sql += " ORDER BY id DESC LIMIT ?"
        params.append(min(max(limit, 1), 500))
        rows = c.execute(sql, params).fetchall()
        c.close()
        cols = ["id", "ts", "category", "rating", "content", "feature", "page", "contact", "status"]
        return {"total": len(rows), "items": [dict(zip(cols, r)) for r in rows]}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ===== 渠道接入配置（WebUI 设置页 → 个人市场切入点）=====
# 各渠道适配器在 channels/ 已实现，配置经 data/channels.json 持久化、路由工厂优先读取。
CHANNEL_DEFS = [
    {
        "key": "telegram", "name": "Telegram", "desc": "机器人渠道，填 Bot Token 即可启用。",
        "doc": "向 @BotFather 申请 Bot 后取 Token 填入。",
        "fields": [{"key": "TELEGRAM_BOT_TOKEN", "label": "Bot Token", "secret": True}],
    },
    {
        "key": "discord", "name": "Discord", "desc": "服务器机器人渠道。",
        "doc": "开发者后台创建应用→Bot→复制 Token。",
        "fields": [{"key": "DISCORD_BOT_TOKEN", "label": "Bot Token", "secret": True}],
    },
    {
        "key": "webhook", "name": "Webhook", "desc": "始终可用，第三方系统可 POST 调用。",
        "doc": "向 /api/webhook/{name} 发送 JSON {text,user_id,chat_id}。",
        "fields": [{"key": "WEBHOOK_SECRET", "label": "回调密钥（可选）", "secret": True}],
    },
    {
        "key": "wecom", "name": "企业微信", "desc": "国内主战场：企业自建应用。",
        "doc": "后台填 corp_id/agent_id/secret；回调 URL 配到企微后台。",
        "fields": [
            {"key": "WECHAT_WORK_CORP_ID", "label": "Corp ID", "secret": False},
            {"key": "WECHAT_WORK_AGENT_ID", "label": "Agent ID", "secret": False},
            {"key": "WECHAT_WORK_SECRET", "label": "Secret", "secret": True},
            {"key": "WECHAT_WORK_TOKEN", "label": "回调 Token（可选）", "secret": True},
            {"key": "WECHAT_WORK_AES_KEY", "label": "AES Key（可选）", "secret": True},
        ],
    },
    {
        "key": "feishu", "name": "飞书", "desc": "国内主战场：企业自建应用。",
        "doc": "开发者后台创建应用，取 App ID/App Secret；回调 URL 配到飞书后台。",
        "fields": [
            {"key": "FEISHU_APP_ID", "label": "App ID", "secret": False},
            {"key": "FEISHU_APP_SECRET", "label": "App Secret", "secret": True},
            {"key": "FEISHU_VERIFY_TOKEN", "label": "Verify Token（可选）", "secret": True},
            {"key": "FEISHU_ENCRYPT_KEY", "label": "Encrypt Key（可选）", "secret": True},
        ],
    },
    {
        "key": "dingtalk", "name": "钉钉", "desc": "国内主战场：企业内部应用。",
        "doc": "开发者后台创建应用，取 AppKey/AppSecret；回调 URL 配到钉钉后台。",
        "fields": [
            {"key": "DINGTALK_APP_KEY", "label": "App Key", "secret": False},
            {"key": "DINGTALK_APP_SECRET", "label": "App Secret", "secret": True},
            {"key": "DINGTALK_AGENT_ID", "label": "Agent ID", "secret": False},
            {"key": "DINGTALK_TOKEN", "label": "回调 Token（可选）", "secret": True},
            {"key": "DINGTALK_AES_KEY", "label": "AES Key（可选）", "secret": True},
        ],
    },
]


@app.get("/api/config/channels")
async def get_channels_config(_=Depends(verify_api_key)):
    """返回各渠道配置（密钥掩码）与运行态启用情况。"""
    try:
        cfg = config.load_channels_config()
        enabled = set()
        try:
            from channels.router import channel_router
            enabled = {c["name"] for c in channel_router.get_status() if c.get("enabled")}
        except Exception:
            pass
        out = []
        for ch in CHANNEL_DEFS:
            saved = cfg.get(ch["key"], {}) or {}
            fields = []
            for f in ch["fields"]:
                raw = saved.get(f["key"]) or getattr(config, f["key"], "")
                is_set = bool(raw)
                fields.append({
                    "key": f["key"],
                    "label": f["label"],
                    "secret": f["secret"],
                    "set": is_set,
                    # 密钥不回传明文；非密钥回传真实值便于查看/微调
                    "value": "••••••••" if (f["secret"] and is_set) else (raw if not f["secret"] else ""),
                })
            out.append({
                "key": ch["key"],
                "name": ch["name"],
                "desc": ch["desc"],
                "doc": ch["doc"],
                "enabled": ch["key"] in enabled,
                "fields": fields,
            })
        return {"channels": out}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/config/channels")
async def save_channels_config_endpoint(payload: dict, _=Depends(verify_api_key)):
    """保存渠道配置（合并写入 data/channels.json，并热重载路由）。"""
    try:
        incoming = payload.get("channels", {})
        if not isinstance(incoming, dict):
            raise HTTPException(status_code=422, detail="channels 必须为对象")
        existing = config.load_channels_config()
        for k, v in incoming.items():
            if not isinstance(v, dict):
                continue
            # 仅保留非空字段；空字符串视为「不修改/清空」
            cleaned = {fk: fv for fk, fv in v.items() if fv not in (None, "")}
            if cleaned:
                existing[k] = {**existing.get(k, {}), **cleaned}
            elif k in existing:
                existing.pop(k, None)
        config.save_channels_config(existing)
        # 热重载：停掉旧渠道、按新配置重启（无需整进程重启）
        try:
            from channels.router import channel_router
            await channel_router.reload()
        except Exception as e:
            print(f"[channels] reload skipped: {e}")
        return {"status": "ok", "saved": list(incoming.keys())}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/channels/reload")
async def reload_channels(_=Depends(verify_api_key)):
    """手动热重载渠道（改完回调/Token 后触发）。"""
    try:
        from channels.router import channel_router
        await channel_router.reload()
        return {"status": "ok", "active": [c["name"] for c in channel_router.get_status() if c.get("enabled")]}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))



# --- 自定义轻量限流中间件（兜底全端点；slowapi 装饰器在该环境未触发）---
import time
from collections import defaultdict, deque
from fastapi import Request
from fastapi.responses import JSONResponse

_rl_hits = defaultdict(deque)

@app.middleware("http")
async def rate_limit_middleware(request: Request, call_next):
    path = request.url.path
    # 豁免：健康检查、静态资源、会话管理(用户高频操作)、技能查询
    if path in ("/health", "/metrics", "/") or path.startswith("/static") or path.startswith("/api/health")        or path.startswith("/api/sessions") or path.startswith("/api/skills") or path.startswith("/api/memory"):
        return await call_next(request)
    client = request.client.host if request.client else "unknown"
    now = time.time()
    limit = 10 if path == "/api/auth/login" else 60
    dq = _rl_hits[client]
    while dq and dq[0] <= now - 60:
        dq.popleft()
    if len(dq) >= limit:
        retry = int(60 - (now - dq[0])) + 1
        return JSONResponse(status_code=429, content={"error": "rate_limit_exceeded", "retry_after": retry})
    dq.append(now)
    return await call_next(request)


# --- Provider protocol switching (OpenAI-compatible vs Anthropic) ---
class ProviderProtocolRequest(BaseModel):
    protocol: str  # "openai" | "anthropic"


@app.post("/api/config/provider/{provider_name}/protocol")
async def set_provider_protocol(provider_name: str, req: ProviderProtocolRequest, _=Depends(verify_api_key)):
    """Switch a provider between its OpenAI-compatible and Anthropic-compatible endpoint."""
    from providers.registry import get as _get_provider
    from core.user_config import set_provider_override
    p = _get_provider(provider_name)
    if not p:
        raise HTTPException(status_code=404, detail=f"Provider '{provider_name}' not found")
    proto = (req.protocol or "").strip().lower()
    if proto == "anthropic":
        aurl = getattr(p, "anthropic_base_url", "") or ""
        if not aurl:
            raise HTTPException(status_code=400, detail=f"'{p.display_name}' 未提供 Anthropic 兼容端点")
        set_provider_override(provider_name, {"base_url": aurl})
    elif proto == "openai":
        set_provider_override(provider_name, {})
    else:
        raise HTTPException(status_code=400, detail="protocol 必须是 openai 或 anthropic")
    return {"ok": True, "provider": provider_name, "protocol": proto, "base_url": p.resolve_base_url()}


class ProviderBaseUrlRequest(BaseModel):
    base_url: str = ""


@app.post("/api/config/provider/{provider_name}/base-url")
async def set_provider_base_url(provider_name: str, req: ProviderBaseUrlRequest, _=Depends(verify_api_key)):
    """Persist a custom OpenAI-compatible base_url for a provider (e.g. point the
    OpenAI-compatible provider at SiliconFlow / a local gateway / any compatible service)."""
    from providers.registry import get as _get_provider
    from core.user_config import set_provider_override
    p = _get_provider(provider_name)
    if not p:
        raise HTTPException(status_code=404, detail=f"Provider '{provider_name}' not found")
    url = (req.base_url or "").strip()
    set_provider_override(provider_name, {"base_url": url} if url else {})
    return {"ok": True, "provider": provider_name, "base_url": p.resolve_base_url()}


@app.get("/api/files/{file_id}")
async def get_file(file_id: str, _=Depends(verify_api_key)):
    from tools.file_hub import get_file_meta
    from fastapi.responses import FileResponse
    meta = get_file_meta(file_id)
    if not meta:
        raise HTTPException(status_code=404, detail="file not found")
    return FileResponse(meta["path"], media_type=meta["mime"], filename=meta["name"])


# --- STT 模型后台预热 ---
@app.on_event("startup")
async def _lumu_stt_warmup():
    import threading
    def _w():
        try:
            from tools import stt_local
            if stt_local.is_available():
                stt_local.warmup()
        except Exception:
            pass
    threading.Thread(target=_w, daemon=True).start()


@app.get("/api/usage")
async def api_usage(hours: int = 0, _=Depends(verify_api_key)):
    """返回真实 token 用量汇总与按天分布（来自 tracing 记录，绝不编造）。"""
    tm = get_tracer()
    h = hours if hours and hours > 0 else 24 * 400
    summary = tm.get_cost_summary(h) or {}
    by_day = tm.get_cost_by_day(400)
    return {"summary": summary, "by_day": by_day}


# ── Capability introspection (appended by patch) ──
import os
@app.get("/api/capabilities")
async def api_capabilities(_=Depends(verify_api_key)):
    """Live capability manifest: tools, toolsets, skill packs, saved skills,
    configured providers/models, exposure policy, and backend API routes.
    Lets external clients (and the WebUI) see the agent's full surface."""
    try:
        tools = [t.name for t in agent.tools.list_tools()]
        toolsets = agent.tools.list_toolsets()
        from skills.skill_packs import scan_packs
        packs = [
            {"name": p.get("name"), "description": p.get("description"),
             "always": p.get("always"), "triggers": p.get("triggers")}
            for p in scan_packs()
        ]
        saved = agent.skills.list_all()
        from core.user_config import load_config
        cfg = load_config()
        providers = {
            pid: {"configured": bool(pc.get("api_key")),
                  "enabled_models": pc.get("enabled_models", [])}
            for pid, pc in (cfg.get("providers") or {}).items()
        }
        from tools.exposure import exposure_policy
        policy = exposure_policy()
        routes = sorted({getattr(r, "path", "") for r in app.routes if getattr(r, "path", "")})
        return {
            "ok": True,
            "agent_home": os.getenv("AGENT_HOME", str(Path(__file__).resolve().parent.parent)),
            "tool_count": len(tools),
            "tools": tools,
            "toolsets": toolsets,
            "skill_packs": packs,
            "saved_skills": saved,
            "providers": providers,
            "exposure_policy": policy,
            "api_routes": routes,
        }
    except Exception as e:
        return {"ok": False, "error": str(e)}


# ── 访问守卫中间件：让「点一次带 token 的链接」真正生效 ──────────────
@app.middleware("http")
async def lumu_access_middleware(request: Request, call_next):
    """辅助层，不承担放行/拒绝职责（那是 verify_api_key 的事）。

    - 带正确 ?token= 访问 → 种下 cookie，此后同源请求（含 <img>、EventSource）自动携带
    - 对外暴露且未授权打开首页 → 返回引导页，而不是让前端满屏 401
    """
    try:
        from core.access_guard import (
            auth_disabled, request_is_exposed, token_matches, unauthorized_page,
        )
    except Exception:
        return await call_next(request)

    plant = None
    try:
        if not auth_disabled() and request_is_exposed(request):
            ok = token_matches(request, request.headers.get("authorization", ""))
            if ok and request.query_params.get("token"):
                plant = request.query_params.get("token")
            if not ok and request.url.path == "/":
                return HTMLResponse(unauthorized_page(), status_code=401)
    except Exception:
        pass

    response = await call_next(request)

    if plant:
        response.set_cookie(
            "lumu_token", plant,
            max_age=31536000, httponly=True, samesite="lax", path="/",
        )
    return response
